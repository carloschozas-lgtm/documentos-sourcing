# -*- coding: utf-8 -*-
import re
import io
import zipfile
import streamlit as st
from pathlib import Path

try:
    from docx import Document
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

try:
    import openpyxl
    XLSX_OK = True
except ImportError:
    XLSX_OK = False

BASE_DIR = Path(__file__).parent

TEXTO_ORIG = "SERVICIO DE DISEÑO, APLICACIÓN Y ANÁLISIS DE LA ENCUESTA DE GÉNERO ENAP 2026"
TR_ORIG    = "TR31108756"

TEXTOS_ADICIONALES = [
    "AUDITORÍAS LABORALES PREVENTIVAS PARA EMPRESAS CONTRATISTAS DE ENAP",
]

PATRON_TR = re.compile(r'[A-Z]{2}\d+')

EXCLUIR = {"app_web.py", "doc_renamer.py"}
EXCLUIR_DIRS = {"generados", ".streamlit", ".git", "__pycache__"}


# ── Lógica de reemplazo ───────────────────────────────────────────────────────

def normalizar(txt: str) -> str:
    return (txt
            .replace("“", '"').replace("”", '"')
            .replace("‘", "'").replace("’", "'")
            .replace("«", '"').replace("»", '"')
            .replace("‹", "'").replace("›", "'"))


def _reemplazar_nombre(texto: str, nuevo_nombre: str) -> tuple[str, int]:
    norm = normalizar(texto)
    total = 0
    for texto_buscar in [TEXTO_ORIG] + TEXTOS_ADICIONALES:
        count = norm.count(texto_buscar)
        if count:
            norm = norm.replace(texto_buscar, nuevo_nombre)
            total += count
    return norm, total


def _reemplazar_tr(texto: str, nuevo_tr: str) -> tuple[str, int]:
    nueva, n = PATRON_TR.subn(nuevo_tr, texto)
    return nueva, n


def _fix_parrafo(parr, nuevo_nombre: str, nuevo_tr: str) -> int:
    hits = 0

    nombre_hits = 0
    for run in parr.runs:
        nuevo_txt, n = _reemplazar_nombre(run.text, nuevo_nombre)
        if n:
            run.text = nuevo_txt
            nombre_hits += n
    if nombre_hits == 0 and parr.runs:
        full = "".join(r.text for r in parr.runs)
        nuevo_full, n = _reemplazar_nombre(full, nuevo_nombre)
        if n:
            parr.runs[0].text = nuevo_full
            for run in parr.runs[1:]:
                run.text = ""
            nombre_hits += n
    hits += nombre_hits

    tr_hits = 0
    for run in parr.runs:
        nuevo_txt, n = _reemplazar_tr(run.text, nuevo_tr)
        if n:
            run.text = nuevo_txt
            tr_hits += n
    if tr_hits == 0 and parr.runs:
        full = "".join(r.text for r in parr.runs)
        nuevo_full, n = _reemplazar_tr(full, nuevo_tr)
        if n:
            parr.runs[0].text = nuevo_full
            for run in parr.runs[1:]:
                run.text = ""
            tr_hits += n
    hits += tr_hits

    return hits


def procesar_docx(src: Path, nuevo_nombre: str, nuevo_tr: str) -> tuple[bytes, int]:
    doc = Document(src)
    total = 0

    def walk_parrafos(parrafos):
        nonlocal total
        for p in parrafos:
            total += _fix_parrafo(p, nuevo_nombre, nuevo_tr)

    def walk_tablas(tablas):
        for tabla in tablas:
            for fila in tabla.rows:
                for celda in fila.cells:
                    walk_parrafos(celda.paragraphs)
                    walk_tablas(celda.tables)

    walk_parrafos(doc.paragraphs)
    walk_tablas(doc.tables)

    for seccion in doc.sections:
        for parte in (
            seccion.header, seccion.footer,
            seccion.even_page_header, seccion.even_page_footer,
            seccion.first_page_header, seccion.first_page_footer,
        ):
            if parte is not None:
                walk_parrafos(parte.paragraphs)
                walk_tablas(parte.tables)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), total


def procesar_xlsx(src: Path, nuevo_nombre: str, nuevo_tr: str) -> tuple[bytes, int]:
    wb = openpyxl.load_workbook(src)
    total = 0
    for ws in wb.worksheets:
        for fila in ws.iter_rows():
            for celda in fila:
                if isinstance(celda.value, str):
                    v, n1 = _reemplazar_nombre(celda.value, nuevo_nombre)
                    nuevo_val, n2 = _reemplazar_tr(v, nuevo_tr)
                    n = n1 + n2
                    if n:
                        celda.value = nuevo_val
                        total += n
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue(), total


# ── Interfaz Streamlit ────────────────────────────────────────────────────────

st.set_page_config(
    page_title="Documentos Sourcing",
    page_icon="📄",
    layout="wide",
)

st.title("📄 Documentos Sourcing")
st.caption("Ingresa los datos del nuevo proceso y descarga los documentos actualizados.")
st.divider()

nuevo_nombre = st.text_area(
    "Nuevo Nombre del Requerimiento",
    value=TEXTO_ORIG,
    height=100,
)

nuevo_tr = st.text_input(
    "Nuevo Código TR",
    value=TR_ORIG,
)

st.divider()

if st.button("▶ Generar Documentos", type="primary", use_container_width=True):

    if not nuevo_nombre.strip():
        st.warning("Ingresa el nuevo nombre del requerimiento.")
        st.stop()
    if not nuevo_tr.strip():
        st.warning("Ingresa el nuevo código TR.")
        st.stop()

    archivos = [
        p for p in BASE_DIR.rglob("*")
        if p.suffix.lower() in {".docx", ".xlsx"}
        and p.name not in EXCLUIR
        and not any(d in p.parts for d in EXCLUIR_DIRS)
    ]

    if not archivos:
        st.warning("No se encontraron plantillas en el repositorio.")
        st.stop()

    resultados = []
    errores = []

    with st.spinner(f"Procesando {len(archivos)} archivos..."):
        zip_buf = io.BytesIO()
        with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
            for src in sorted(archivos):
                rel = src.relative_to(BASE_DIR)
                nombre_nuevo = PATRON_TR.sub(nuevo_tr.strip(), rel.name)
                zip_path = str(rel.parent / nombre_nuevo)
                try:
                    if src.suffix.lower() == ".docx":
                        resultado_bytes, n = procesar_docx(src, nuevo_nombre.strip(), nuevo_tr.strip())
                    else:
                        resultado_bytes, n = procesar_xlsx(src, nuevo_nombre.strip(), nuevo_tr.strip())
                    zf.writestr(zip_path, resultado_bytes)
                    resultados.append((str(rel), nombre_nuevo, n))
                except Exception as exc:
                    errores.append((str(rel), str(exc)))

    total_reempl = sum(n for _, _, n in resultados)
    st.success(f"✅ {len(resultados)} archivo(s) generado(s) · {total_reempl} reemplazo(s) en total")

    zip_buf.seek(0)
    st.download_button(
        label="⬇️ Descargar todos los documentos (.zip)",
        data=zip_buf,
        file_name=f"Documentos_{nuevo_tr.strip()}.zip",
        mime="application/zip",
        use_container_width=True,
    )

    with st.expander("Ver detalle por archivo"):
        for rel, nombre_nuevo, n in resultados:
            icono = "✔️" if n else "➖"
            st.write(f"{icono} `{rel}` → `{nombre_nuevo}` — {n} reemplazo{'s' if n != 1 else ''}")
        for rel, err in errores:
            st.error(f"✘ `{rel}` — {err}")
